from fastapi import FastAPI, UploadFile, File, Form, Response
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Dict, Any

# ... your existing imports (read_resume or inline), and analyze import ...
import io
import re
import string
from collections import Counter
from typing import Tuple, List, Dict, Any

# Optional parsers (gracefully handled if missing)
try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None


try:
    import docx  # python-docx
except Exception:
    docx = None

# ------------------------- Utilities & Config -------------------------

# A small, extensible skills/keywords library; add more as you like
SKILL_LIBRARY = {
    # Languages
    "python", "java", "c++", "c", "javascript", "typescript", "go", "rust",
    # Backend / Web
    "fastapi", "flask", "django", "express", "node.js", "nodejs", "rest", "graphql",
    "spring", "spring boot",
    # Databases
    "postgresql", "mysql", "sqlite", "mongodb", "redis",
    # DevOps
    "docker", "kubernetes", "k8s", "aws", "gcp", "azure", "ci/cd", "nginx",
    # Data/ML
    "pandas", "numpy", "scikit-learn", "sklearn", "tensorflow", "pytorch",
    # Tools
    "git", "github", "gitlab", "jira", "linux", "bash", "shell", "vim",
    # Testing/Other
    "pytest", "unittest", "selenium",
    # Cloud/Infra extras
    "terraform", "ansible",
}

SECTION_HEADINGS = [
    "education", "experience", "work experience", "projects", "skills",
    "achievements", "certifications", "publications", "summary", "objective"
]

EMAIL_RE = re.compile(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b")
PHONE_RE = re.compile(r"(?:\+?\d[\s-]?){7,15}")  # permissive for intl formats
URL_RE = re.compile(r"https?://\S+|www\.\S+")
WHITESPACE_RE = re.compile(r"\s+")

def _safe_decode(b: bytes) -> str:
    try:
        return b.decode("utf-8", errors="ignore")
    except Exception:
        try:
            return b.decode("latin-1", errors="ignore")
        except Exception:
            return ""

def _normalize_text(text: str) -> str:
    # Lowercase, collapse spaces; keep punctuation for section detection
    text = text.replace("\x00", " ")
    text = WHITESPACE_RE.sub(" ", text)
    return text

def _words(text: str) -> List[str]:
    # tokenization: split on non-alphanumeric; keep simple
    text = text.lower()
    text = text.translate(str.maketrans({c: " " for c in string.punctuation}))
    return [w for w in text.split() if w]

def _find_snippets(text: str, terms: List[str], window: int = 60) -> Dict[str, List[str]]:
    """Return small context windows for matched terms."""
    out: Dict[str, List[str]] = {}
    lower = text.lower()
    for t in terms:
        t_low = t.lower()
        idx = 0
        hits = []
        while True:
            idx = lower.find(t_low, idx)
            if idx == -1:
                break
            start = max(0, idx - window)
            end = min(len(text), idx + len(t) + window)
            snippet = text[start:end].strip()
            hits.append(snippet)
            idx += len(t_low)
        if hits:
            out[t] = hits[:3]  # cap snippets per term
    return out

def _extract_keywords_from_jd(jd_text: str) -> List[str]:
    """
    Heuristic: intersect JD words/phrases with a curated library + keep multiword phrases as-is.
    Also pull capitalized acronyms (e.g., AWS, GCP) and common tech tokens.
    """
    jd_norm = _normalize_text(jd_text)
    jd_tokens = set(_words(jd_norm))

    # Include multi-word phrases that exist in SKILL_LIBRARY
    multi_phrases = {s for s in SKILL_LIBRARY if " " in s and all(w in jd_tokens for w in s.split())}

    # Acronyms / shortcaps from JD
    acronyms = set(re.findall(r"\b[A-Z]{2,}\b", jd_text))
    # Normalize acronyms like CI/CD -> cicd
    acronyms_norm = {a.lower().replace("/", "") for a in acronyms}

    # Basic intersection with skill library (single tokens)
    singles = {s for s in SKILL_LIBRARY if " " not in s and s in jd_tokens}

    # Merge
    kws = sorted(singles | multi_phrases | acronyms_norm)
    return kws

# ------------------------- Resume Readers -------------------------

def read_resume(upload_file: UploadFile) -> str:
    """
    Read text from PDF/DOCX/TXT. Falls back gracefully; returns plain text string.
    """
    name = (upload_file.filename or "").lower()
    ctype = (upload_file.content_type or "").lower()

    data = awaitable_read(upload_file)  # read bytes synchronously from UploadFile.file

    # Prefer by extension if present, else use content-type
    if name.endswith(".pdf") or "pdf" in ctype:
        if fitz is None:
            txt = _safe_decode(data)
            return _normalize_text(
                f"[WARN: PyMuPDF not installed; returning raw bytes instead]\n{txt}"
            )
        try:
        # Read PDF bytes with PyMuPDF
            text = ""
            with fitz.open(stream=data, filetype="pdf") as doc:
                for page in doc:
                # Extract text with layout-aware mode
                    text += page.get_text("text") + "\n"

                # Optionally also extract hyperlinks to catch portfolio URLs
                    for link in page.get_links():
                        if "uri" in link:
                            text += f"\n{link['uri']}"
            return _normalize_text(text)
        except Exception as e:
        # Fallback: if PyMuPDF fails, decode bytes anyway
            txt = _safe_decode(data)
            return _normalize_text(f"[WARN: PyMuPDF parse failed: {e}]\n{txt}")


    if name.endswith(".docx") or "officedocument.wordprocessingml.document" in ctype:
        if docx is None:
            txt = _safe_decode(data)
            return _normalize_text(
                f"[WARN: python-docx not installed; returning raw text bytes]\n{txt}"
            )
        try:
            f = io.BytesIO(data)
            d = docx.Document(f)
            text = "\n".join(p.text for p in d.paragraphs)
            # include tables if any
            for table in d.tables:
                for row in table.rows:
                    text += "\n" + " | ".join(cell.text for cell in row.cells)
            return _normalize_text(text)
        except Exception:
            return _normalize_text(_safe_decode(data))

    # Plain text or unknown -> try decode
    return _normalize_text(_safe_decode(data))

def awaitable_read(upload_file: UploadFile) -> bytes:
    """
    UploadFile.file is a SpooledTemporaryFile. We can read its bytes here.
    (Kept separate so you can swap to an async read if you prefer.)
    """
    upload_file.file.seek(0)
    return upload_file.file.read()

# ------------------------- Analyzer -------------------------

def analyze(resume_text: str, jd_text: str) -> Dict[str, Any]:
    """
    Lightweight analyzer:
      - Extract JD keywords (heuristic + skill library)
      - Match presence in resume
      - Compute coverage score
      - Run ATS checks (contact info, sections, links, file hints)
      - Highlights/snippets for matched keywords
      - Simple "semantic" overlap via top n-grams
      - Draft rewritten lines for missing keywords
    """
    # Normalize for search but keep original for snippets
    resume_norm = resume_text.lower()
    jd_norm = jd_text.lower()

    jd_keywords = _extract_keywords_from_jd(jd_text)
    # Also add most frequent content words from JD (helps when library misses)
    jd_tokens = [w for w in _words(jd_norm) if len(w) > 2]
    common_jd = {w for w, c in Counter(jd_tokens).most_common(20) if w not in jd_keywords}
    # Merge (limit growth)
    all_targets = sorted(set(jd_keywords) | set(list(common_jd)[:10]))

    # Match / Missing
    matched, missing = [], []
    for t in all_targets:
        if " " in t:
            present = t in resume_norm
        else:
            # whole-word-ish check
            present = re.search(rf"\b{re.escape(t)}\b", resume_norm) is not None
        (matched if present else missing).append(t)

    # Score (coverage + small bonus for sections and contact)
    base_score = int(round(100 * (len(matched) / max(1, len(all_targets)))))
    ats = _ats_checks(resume_text)
    bonus = 0
    if ats["has_email"] and ats["has_phone"]:
        bonus += 3
    if ats["has_sections"]:
        bonus += 2
    score = max(0, min(100, base_score + bonus))

    # Highlights
    highlights = _find_snippets(resume_text, matched)

    # "Semantic" matches via n-gram overlap (very simple)
    semantic = _semantic_overlap(jd_text, resume_text)

    # Recommendations for missing keywords
    recs, rewrites = _recommendations(missing, jd_text)

    result: Dict[str, Any] = {
        "score": score,
        "matched_keywords": matched,
        "missing_keywords": missing,
        "recommendations": recs,
        "ats_checks": ats,
        "highlights": highlights,
        "semantic_matches": semantic,
        "rewritten_examples": rewrites,
    }
    return result

def _ats_checks(text: str) -> Dict[str, Any]:
    has_email = EMAIL_RE.search(text) is not None
    has_phone = PHONE_RE.search(text) is not None
    has_links = URL_RE.search(text) is not None

    lower = text.lower()
    found_sections = [s for s in SECTION_HEADINGS if s in lower]
    has_sections = len(found_sections) >= 2

    # File hints (helpful if you pass through warning banners from read_resume)
    file_warnings = "[WARN:" in text

    return {
        "has_email": has_email,
        "has_phone": has_phone,
        "has_links": has_links,
        "has_sections": has_sections,
        "detected_sections": found_sections,
        "file_parser_warning": file_warnings,
    }

def _semantic_overlap(src: str, tgt: str, n: int = 3, top_k: int = 10) -> List[Dict[str, Any]]:
    """
    Crude semantic signal: top N-grams from JD that also occur in resume.
    """
    def ngrams(words: List[str], n: int) -> List[str]:
        return [" ".join(words[i:i+n]) for i in range(len(words) - n + 1)]

    sw = set({
        "and","or","the","is","to","of","a","in","on","for","with","by","as","at","an","be",
        "this","that","it","from","are","you","your","our","we","their","they","will","can"
    })

    s_words = [w for w in _words(src) if w not in sw]
    t_lower = tgt.lower()

    grams = ngrams(s_words, n)
    counts = []
    for g in grams:
        present = g in t_lower
        if present:
            counts.append(g)

    # Rank by frequency in JD
    freq = Counter(grams)
    top = []
    for g, c in freq.most_common():
        if g in counts:
            top.append({"ngram": g, "present_in_resume": True, "jd_frequency": c})
        if len(top) >= top_k:
            break
    return top

def _recommendations(missing: List[str], jd_text: str) -> Tuple[List[str], List[str]]:
    recs: List[str] = []
    rewrites: List[str] = []

    if missing:
        recs.append(
            f"Consider adding concrete bullet points that demonstrate {', '.join(missing[:6])} "
            "with metrics (impact %, latency ms, cost $, users served)."
        )

    # Generic structure suggestions
    recs += [
        "Add a Skills section grouping programming languages, frameworks, databases, and tools.",
        "Quantify achievements (e.g., 'cut API latency by 35%' or 'handled 50k req/day').",
        "Use standard headings: Education, Experience, Projects, Skills.",
        "Link to GitHub/portfolio and ensure keywords from the JD appear naturally in bullets."
    ]

    # Example rewrites for top missing items
    for t in missing[:5]:
        rewrites.append(f"Implemented {t} in production APIs; improved reliability and reduced p99 latency.")
    return recs, rewrites



app = FastAPI(title="NLP Resume Analyzer", version="2.0")


# (NEW) make /analyze and /analyze/ behave the same
app.router.redirect_slashes = True

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],          # must include OPTIONS
    allow_headers=["*"],
)

# (NEW) Handle preflight cleanly if the browser sends it
@app.options("/analyze")
def options_analyze():
    return Response(status_code=204)

# (optional) also handle trailing slash explicitly
@app.options("/analyze/")
def options_analyze_slash():
    return Response(status_code=204)

# ---------------- your existing routes ----------------

class AnalysisResponse(BaseModel):
    score: int
    matched_keywords: List[str]
    missing_keywords: List[str]
    recommendations: List[str]
    ats_checks: Dict[str, Any]
    highlights: Dict[str, Any]
    semantic_matches: List[Dict[str, Any]]
    rewritten_examples: List[str]

@app.post("/analyze", response_model=AnalysisResponse)
async def analyze_resume(
    resume: UploadFile = File(...),
    jd: str = Form(...)
):
    resume_text = read_resume(resume)
    result = analyze(resume_text, jd)
    return AnalysisResponse(**result)

# (optional) mirror POST on slash route to be extra-forgiving
@app.post("/analyze/")
async def analyze_resume_slash(
    resume: UploadFile = File(...),
    jd: str = Form(...)
):
    resume_text = read_resume(resume)
    result = analyze(resume_text, jd)
    return result

@app.get("/")
def root():
    return {"ok": True, "message": "NLP Resume Analyzer API. POST /analyze with resume + jd."}